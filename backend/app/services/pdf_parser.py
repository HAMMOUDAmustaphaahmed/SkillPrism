import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> Optional[str]:
    """Extract text from PDF bytes using pdfplumber."""
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages_text = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages_text.append(text)
            return "\n".join(pages_text)
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        return None


def extract_text_from_docx(file_bytes: bytes) -> Optional[str]:
    """Extract text from DOCX bytes using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return None


def extract_text_from_txt(file_bytes: bytes) -> Optional[str]:
    """Extract text from plain text bytes."""
    try:
        return file_bytes.decode("utf-8", errors="ignore")
    except Exception as e:
        logger.error(f"TXT extraction error: {e}")
        return None


def parse_file(file_bytes: bytes, filename: str) -> Optional[str]:
    """Route file to correct parser based on extension."""
    name_lower = filename.lower()
    if name_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif name_lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    elif name_lower.endswith(".txt"):
        return extract_text_from_txt(file_bytes)
    else:
        return extract_text_from_txt(file_bytes)
